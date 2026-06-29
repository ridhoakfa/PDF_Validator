import io
from docx import Document
from docx.shared import Pt, Inches

def extract_text_from_docx(file_bytes):
    """Ekstrak teks dari DOCX"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        return f"Error: {e}"

def get_docx_info(file_bytes):
    """Dapatkan informasi format dari DOCX: font, ukuran, spasi, margin, jumlah kata"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        
        fonts = set()
        font_sizes = set()
        line_spacings = []
        paragraphs = 0
        total_words = 0
        
        for para in doc.paragraphs:
            paragraphs += 1
            # Hitung kata
            words = len(para.text.split())
            total_words += words
            
            # Spasi baris
            if para.paragraph_format.line_spacing:
                if isinstance(para.paragraph_format.line_spacing, float):
                    line_spacings.append(para.paragraph_format.line_spacing)
                elif isinstance(para.paragraph_format.line_spacing, Pt):
                    line_spacings.append(para.paragraph_format.line_spacing.pt)
            
            # Font dan ukuran
            for run in para.runs:
                if run.font.name:
                    fonts.add(run.font.name)
                if run.font.size:
                    font_sizes.add(run.font.size.pt)
        
        # Margin dokumen (section)
        margins = []
        for section in doc.sections:
            margins.append({
                "top": section.top_margin.inches if section.top_margin else 0,
                "bottom": section.bottom_margin.inches if section.bottom_margin else 0,
                "left": section.left_margin.inches if section.left_margin else 0,
                "right": section.right_margin.inches if section.right_margin else 0
            })
        
        # Rata kanan-kiri? Cek alignment
        align_justify = any(para.alignment == 3 for para in doc.paragraphs)  # WD_ALIGN_PARAGRAPH.JUSTIFY = 3
        
        return {
            "total_words": total_words,
            "paragraphs": paragraphs,
            "fonts": list(fonts),
            "font_sizes": list(font_sizes),
            "line_spacings": line_spacings,
            "margins": margins,
            "align_justify": align_justify,
            "is_times_new_roman": any("Times" in f or "TimesNewRoman" in f for f in fonts),
            "font_size_12": any(11.5 <= s <= 12.5 for s in font_sizes),
            "spacing_1_5": any(1.4 <= s <= 1.6 for s in line_spacings)
        }
    except Exception as e:
        return {"error": str(e)}